#!/usr/bin/env python3
# -*- coding: utf-8 -*-
r"""
Parse Word catalogs from C:\Users\maria\Pictures\Inventory Unistyles images
into structured JSON for Strapi seed.

Rules:
  - Red highlight on any cell -> product discontinued, dropped
  - Bras (BH catalog) grouped by REF: 1 parent product, N variants (color x size)
  - Other catalogs: flat products
  - Prices normalized: "$19", "Xcg 30", "XCG 75" -> 19, 30, 75 (integer XCG)
"""

import json
import re
import sys
from pathlib import Path
from collections import defaultdict

import docx
from docx.enum.text import WD_COLOR_INDEX

SOURCE_DIR = Path("/mnt/c/Users/maria/Pictures/Inventory Unistyles images")
OUTPUT_DIR = Path("/tmp/inv-parsed")
OUTPUT_DIR.mkdir(exist_ok=True)


def cell_text(cell):
    return "\n".join(p.text for p in cell.paragraphs).strip()


def cell_is_red(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            if run.font.highlight_color == WD_COLOR_INDEX.RED:
                return True
    return False


def row_is_red(row):
    return any(cell_is_red(c) for c in row.cells)


def normalize_price(text):
    if not text:
        return None
    nums = re.findall(r"\d+(?:[.,]\d+)?", text.replace(",", "."))
    if not nums:
        return None
    return float(nums[0])


def normalize_stock(text):
    if not text:
        return 0
    m = re.search(r"\d+", text)
    return int(m.group()) if m else 0


def parse_table(tbl, schema):
    """Generic table parser. schema = list of column names matching header."""
    rows = list(tbl.rows)
    if not rows:
        return []

    # Find header row (contains "Product Name" or "SKU")
    header_idx = 0
    for i, row in enumerate(rows[:3]):
        joined = " ".join(cell_text(c).lower() for c in row.cells)
        if "product name" in joined or "sku" in joined:
            header_idx = i
            break

    header = [cell_text(c).lower().strip().replace("**", "") for c in rows[header_idx].cells]

    items = []
    for row in rows[header_idx + 1:]:
        if len(row.cells) < 2:
            continue
        record = {"_red": row_is_red(row)}
        for i, cell in enumerate(row.cells):
            if i >= len(header):
                break
            key = header[i] or f"col{i}"
            record[key] = cell_text(cell)
        # Skip empty rows (no product name)
        name_field = record.get("product name", "") or record.get("name", "")
        if not name_field.strip():
            continue
        items.append(record)
    return items


def parse_catalog(path, category_slug):
    doc = docx.Document(str(path))
    all_items = []
    for tbl in doc.tables:
        items = parse_table(tbl, [])
        all_items.extend(items)
    return all_items


def build_bras(raw_items):
    """Bras: group by reference number, build parent + variants array."""
    by_ref = defaultdict(list)
    for it in raw_items:
        if it.get("_red"):
            continue
        ref = (it.get("reference") or "").strip()
        if not ref:
            # Try to extract REF from name
            name = it.get("product name", "")
            m = re.search(r"REF\s*0*(\d+)", name, re.I)
            if m:
                ref = m.group(1)
        if not ref:
            continue
        by_ref[ref].append(it)

    products = []
    for sort_idx, (ref, group) in enumerate(sorted(by_ref.items()), start=1):
        first = group[0]
        # Strip "REF NNN" and color/size from name to get clean parent name
        parent_name = re.sub(r"\s*(BEIGE|BLANCO|NEGRO|BLACK|WHITE|GRIS|ROSA|CREMA|LILA|MORADO)\s*",
                             "", first.get("product name", ""), flags=re.I)
        parent_name = re.sub(r"\s*REF\s*0*\d+\s*", "", parent_name, flags=re.I)
        parent_name = re.sub(r"\s*\d+\s*[A-D]\s*", "", parent_name)
        parent_name = re.sub(r"\s+", " ", parent_name).strip()
        if not parent_name:
            parent_name = f"Leonisa Bra REF {ref}"

        first_price = normalize_price(first.get("price", ""))

        variants = []
        for v in group:
            variants.append({
                "sku": v.get("sku", "").strip(),
                "color": v.get("color", "").strip(),
                "size": v.get("size", "").strip(),
                "price": normalize_price(v.get("price", "")),
                "stockQuantity": normalize_stock(v.get("stock", "")),
                "inStock": normalize_stock(v.get("stock", "")) > 0,
            })

        products.append({
            "id": f"BRA-PARENT-{ref}",
            "ref": ref,
            "name": parent_name,
            "brand": first.get("brand", "Leonisa").strip() or "Leonisa",
            "price": first_price,
            "description": first.get("description", "").strip(),
            "image": f"/images/bra/{ref}/main.png",
            "variants": variants,
            "totalStock": sum(v["stockQuantity"] for v in variants),
            "inStock": any(v["inStock"] for v in variants),
            "sortOrder": sort_idx,
        })
    return products


DEFAULT_PRICES = {
    "perfume": 19,
    "cremas": 25,
    "bloqueador": 75,
    "limpieza-facial": 35,
}


def build_flat(raw_items, sku_prefix, category_slug, auto_sku_start=None):
    """Flat products: each row is its own product. Drops red rows."""
    products = []
    auto_sku_counter = auto_sku_start or 100
    sort_idx = 1
    fallback_price = DEFAULT_PRICES.get(category_slug, 0)
    for it in raw_items:
        if it.get("_red"):
            continue
        sku = (it.get("sku") or "").strip()
        if not sku and auto_sku_start is not None:
            sku = f"{sku_prefix}-{auto_sku_counter:03d}"
            auto_sku_counter += 1
        if not sku:
            continue
        name = it.get("product name", "").strip()
        if not name:
            continue

        price = normalize_price(it.get("price", ""))
        if price is None:
            price = fallback_price
            print(f"  [warn] {sku} '{name}' has no price in catalog — defaulted to {price}", file=sys.stderr)

        products.append({
            "id": sku,
            "ref": sku,
            "name": name,
            "brand": (it.get("brand") or "").strip(),
            "price": price,
            "description": (it.get("description") or "").strip(),
            "stockQuantity": normalize_stock(it.get("stock", "")),
            "inStock": normalize_stock(it.get("stock", "")) > 0,
            "sortOrder": sort_idx,
            # image filled in later by mapping step
            "image": None,
        })
        sort_idx += 1
    return products


def main():
    catalogs = {
        "bras": ("Copia de CATALOGO_BH_FINAL.docx", "BRA", build_bras, None),
        "perfume": ("CATALOGO_COLONIAS_UPDATED.docx", "COL", "flat", None),
        "cremas": ("CATALOGO_CREMAS_UPDATED.docx", "CRM", "flat", None),
        "bloqueador": ("CATALOGO_BLOQUEADOR_UPDATED.docx", "BLQ", "flat", 12),
        "limpieza-facial": ("CATALOGO_LIMPIEZA_FACIAL_UPDATED.docx", "LF", "flat", None),
    }

    summary = []
    output = {}

    for slug, (filename, prefix, builder, auto_start) in catalogs.items():
        path = SOURCE_DIR / filename
        if not path.exists():
            print(f"  MISSING: {filename}", file=sys.stderr)
            continue

        raw = parse_catalog(path, slug)
        red_count = sum(1 for r in raw if r.get("_red"))

        if builder == build_bras:
            products = build_bras(raw)
            variant_total = sum(len(p["variants"]) for p in products)
            summary.append((slug, len(raw), red_count, len(products), variant_total))
        else:
            products = build_flat(raw, prefix, slug, auto_start)
            summary.append((slug, len(raw), red_count, len(products), 0))

        output[slug] = products

    # Print summary
    print("\n" + "=" * 78)
    print(f"{'Category':<20} {'Total rows':>12} {'Red (drop)':>12} {'Products':>10} {'Variants':>10}")
    print("=" * 78)
    for slug, total, red, prods, variants in summary:
        print(f"{slug:<20} {total:>12} {red:>12} {prods:>10} {variants:>10}")
    print("=" * 78)

    # Write per-category JSON
    for slug, products in output.items():
        out_path = OUTPUT_DIR / f"{slug}.json"
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(products, f, indent=2, ensure_ascii=False)
        print(f"  wrote {out_path} ({len(products)} products)")

    # Write combined output too
    combined_path = OUTPUT_DIR / "all-parsed.json"
    with open(combined_path, "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, ensure_ascii=False)
    print(f"  wrote {combined_path}")


if __name__ == "__main__":
    main()
